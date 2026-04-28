"""Build the AI Identity Discovery Document as an editable Word (.docx) file.

Clients fill it out in Word / Google Docs / Pages before the discovery call.
Jeff's Notes section at the end is clearly marked as internal.

Run: python3 build_discovery_doc_word.py
Output: ../discovery-doc.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "discovery-doc.docx"

# ── Brand colours ──────────────────────────────────────────────────────────
NAVY = RGBColor(0x0B, 0x19, 0x33)
AMBER = RGBColor(0xE0, 0x9F, 0x3E)
INK = RGBColor(0x1A, 0x1F, 0x36)
GREY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GREY = RGBColor(0x9C, 0xA3, 0xAF)
LIGHT_BG = RGBColor(0xF5, 0xF7, 0xFA)
FIELD_BG = RGBColor(0xFA, 0xFB, 0xFC)
NOTES_BG = RGBColor(0xEE, 0xF3, 0xFB)
BLUE = RGBColor(0x4F, 0x8F, 0xE5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LINE = RGBColor(0xE2, 0xE5, 0xEB)

doc = Document()


# ── Page setup ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)


# ── Default style tweaks ───────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = INK


# ── XML helpers ────────────────────────────────────────────────────────────


def rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()  # noqa: N806
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, *, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()  # noqa: N806
    tcBorders = OxmlElement("w:tcBorders")  # noqa: N806
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "E2E5EB"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def no_border(cell):
    for side in ["top", "bottom", "left", "right"]:
        set_cell_border(cell, **{side: {"val": "none", "sz": 0, "color": "FFFFFF"}})


def set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()  # noqa: N806
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def add_run(para, text, *, bold=False, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def spacer(height_pt=4):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=0)
    run = p.add_run()
    run.font.size = Pt(height_pt)


# ── Hero header table ──────────────────────────────────────────────────────

hero = doc.add_table(rows=1, cols=2)
hero.alignment = WD_TABLE_ALIGNMENT.LEFT
hero.style = "Table Grid"

left_cell = hero.cell(0, 0)
right_cell = hero.cell(0, 1)

# Amber left bar — narrow column
hero.columns[0].width = Inches(0.18)
hero.columns[1].width = Inches(6.29)

set_cell_bg(left_cell, AMBER)
no_border(left_cell)
left_cell.paragraphs[0].add_run()

set_cell_bg(right_cell, NAVY)
no_border(right_cell)
right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p1 = right_cell.paragraphs[0]
set_para_spacing(p1, before=60, after=0)
add_run(p1, "AI IDENTITY", bold=True, size=9, color=AMBER)

p2 = right_cell.add_paragraph()
set_para_spacing(p2, before=20, after=0)
add_run(p2, "Discovery Document", bold=True, size=22, color=WHITE)

p3 = right_cell.add_paragraph()
set_para_spacing(p3, before=30, after=0)
add_run(
    p3,
    "Please fill out what you can before our call. We'll review it together and build on it.",
    size=9.5,
    color=RGBColor(0xC9, 0xD4, 0xE6),
)

p4 = right_cell.add_paragraph()
set_para_spacing(p4, before=6, after=60)
add_run(p4, "jeff@ai-identity.co  ·  ai-identity.co", size=8.5, color=RGBColor(0x9D, 0xA9, 0xC1))

spacer(6)


# ── Section helper ─────────────────────────────────────────────────────────


def section_header(num: str, title: str, subtitle: str = ""):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(0.1)
    tbl.columns[1].width = Inches(3.8)
    tbl.columns[2].width = Inches(2.57)

    amber_cell = tbl.cell(0, 0)
    set_cell_bg(amber_cell, AMBER)
    no_border(amber_cell)
    amber_cell.paragraphs[0].add_run()

    title_cell = tbl.cell(0, 1)
    set_cell_bg(title_cell, NAVY)
    no_border(title_cell)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = title_cell.paragraphs[0]
    set_para_spacing(p, before=50, after=50)
    add_run(p, f"{num}  ", bold=True, size=8, color=AMBER)
    add_run(p, title.upper(), bold=True, size=10, color=WHITE)

    sub_cell = tbl.cell(0, 2)
    set_cell_bg(sub_cell, NAVY)
    no_border(sub_cell)
    sub_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p2 = sub_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p2, before=50, after=50)
    if subtitle:
        add_run(p2, subtitle, size=8, color=RGBColor(0x9D, 0xA9, 0xC1))

    spacer(4)


# ── Field helpers ──────────────────────────────────────────────────────────


def field(label: str, hint: str = "", lines: int = 1):
    """Label above a shaded editable area."""
    lp = doc.add_paragraph()
    set_para_spacing(lp, before=30, after=20)
    add_run(lp, label, bold=True, size=8.5, color=INK)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, FIELD_BG)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
        bottom={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
        left={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
        right={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
    )
    p = cell.paragraphs[0]
    set_para_spacing(p, before=40, after=40)
    if hint:
        add_run(p, hint, size=8, color=LIGHT_GREY, italic=True)
    # blank lines for writing space
    for _ in range(lines - 1):
        np = cell.add_paragraph()
        set_para_spacing(np, before=0, after=0)
        add_run(np, "", size=10)
    spacer(4)


def field_pair(label_a: str, label_b: str, hint_a: str = "", hint_b: str = ""):
    """Two fields side by side."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(3.1)
    tbl.columns[1].width = Inches(3.37)

    for col, (label, hint) in enumerate([(label_a, hint_a), (label_b, hint_b)]):
        label_cell = tbl.cell(0, col)
        no_border(label_cell)
        set_cell_bg(label_cell, WHITE)
        lp = label_cell.paragraphs[0]
        set_para_spacing(lp, before=30, after=10)
        add_run(lp, label, bold=True, size=8.5, color=INK)

        input_cell = tbl.cell(1, col)
        set_cell_bg(input_cell, FIELD_BG)
        set_cell_border(
            input_cell,
            top={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
            bottom={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
            left={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
            right={"val": "single", "sz": 4, "color": rgb_hex(LINE)},
        )
        ip = input_cell.paragraphs[0]
        set_para_spacing(ip, before=40, after=40)
        if hint:
            add_run(ip, hint, size=8, color=LIGHT_GREY, italic=True)

    spacer(4)


def checkbox_list(items: list[str], cols: int = 2):
    """Checkbox grid using a table."""
    rows = -(-len(items) // cols)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, item in enumerate(items):
        r, col = divmod(i, cols)
        cell = tbl.cell(r, col)
        no_border(cell)
        p = cell.paragraphs[0]
        set_para_spacing(p, before=30, after=30)
        add_run(p, "☐  ", bold=False, size=11, color=INK)  # ☐
        add_run(p, item, size=9, color=INK)

    # fill any empty trailing cells
    total_cells = rows * cols
    for i in range(len(items), total_cells):
        r, col = divmod(i, cols)
        cell = tbl.cell(r, col)
        no_border(cell)

    spacer(4)


def notes_block(label: str, lines: int = 6):
    """Blue-tinted notes area for Jeff."""
    lp = doc.add_paragraph()
    set_para_spacing(lp, before=30, after=20)
    add_run(lp, label, bold=True, size=8.5, color=BLUE)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(0.1)
    tbl.columns[1].width = Inches(6.37)

    bar = tbl.cell(0, 0)
    set_cell_bg(bar, BLUE)
    no_border(bar)
    bar.paragraphs[0].add_run()

    content = tbl.cell(0, 1)
    set_cell_bg(content, NOTES_BG)
    set_cell_border(
        content,
        top={"val": "single", "sz": 4, "color": rgb_hex(BLUE)},
        bottom={"val": "single", "sz": 4, "color": rgb_hex(BLUE)},
        right={"val": "single", "sz": 4, "color": rgb_hex(BLUE)},
        left={"val": "none", "sz": 0, "color": "FFFFFF"},
    )
    p = content.paragraphs[0]
    set_para_spacing(p, before=30, after=0)
    add_run(
        p,
        "For internal use — observations, clarifications, and next steps from the call",
        size=8,
        color=GREY,
        italic=True,
    )
    for _ in range(lines):
        np = content.add_paragraph()
        set_para_spacing(np, before=0, after=0)
        add_run(np, "", size=11)
    lp2 = content.add_paragraph()
    set_para_spacing(lp2, before=0, after=30)

    spacer(6)


# ══════════════════════════════════════════════════════════════════════════
# SECTION 01 — Company & Contact
# ══════════════════════════════════════════════════════════════════════════

section_header("01", "Company & Contact", "Fill out before the call")
field_pair("Company name", "Website / LinkedIn")
field_pair("Your name & title", "Best email for follow-up")
field_pair(
    "Other attendees on your side",
    "Preferred call timezone / slot",
    hint_a="Name, title — list everyone joining",
    hint_b="e.g. US-East, afternoons",
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 02 — Agent Inventory
# ══════════════════════════════════════════════════════════════════════════

section_header("02", "Agent Inventory", "Tell us what you're running or building")
field(
    "What agents do you have in production or in development?",
    hint="e.g. customer-support agent, internal data-retrieval agent, code-gen pipeline — list each briefly",
    lines=3,
)
field_pair(
    "Approx. requests / month (total across agents)",
    "Number of distinct agents (today / projected 12 mo)",
    hint_a="e.g. 200K / mo",
    hint_b="e.g. 4 today → 20 in 12 months",
)

lp = doc.add_paragraph()
set_para_spacing(lp, before=30, after=10)
add_run(
    lp,
    "What do these agents have access to? (check all that apply)",
    bold=True,
    size=8.5,
    color=INK,
)
checkbox_list(
    [
        "Customer PII / account data",
        "Internal knowledge base / docs",
        "Financial / billing systems",
        "Code repositories",
        "Third-party APIs (CRM, ERP, etc.)",
        "Email / calendar / communication tools",
        "Healthcare / clinical data",
        "Other (describe below)",
    ],
    cols=2,
)
field("Other access / notes", lines=2)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 03 — Compliance & Regulatory Context
# ══════════════════════════════════════════════════════════════════════════

section_header("03", "Compliance & Regulatory Context")

lp = doc.add_paragraph()
set_para_spacing(lp, before=30, after=10)
add_run(
    lp,
    "Which compliance frameworks apply to your organization? (check all that apply)",
    bold=True,
    size=8.5,
    color=INK,
)
checkbox_list(
    [
        "SOC 2 Type II",
        "ISO 27001",
        "HIPAA",
        "GDPR / EU AI Act",
        "NIST AI RMF",
        "FedRAMP",
        "ITAR / CMMC",
        "None yet — planning ahead",
    ],
    cols=2,
)

field_pair(
    "Are you currently audited / certified for any of these?",
    "Next audit or certification deadline",
    hint_a="Which ones, and by which auditor?",
    hint_b="Approximate date or quarter",
)
field(
    "Do you have an AI governance policy today?",
    hint="Yes / No / In progress — describe briefly if yes",
    lines=2,
)
field(
    "Any specific data-residency or sovereignty requirements?",
    hint="e.g. data must stay in EU, no US-hosted compute, customer contract requirements",
    lines=2,
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 04 — Current Risk & Pain Points
# ══════════════════════════════════════════════════════════════════════════

section_header("04", "Current Risk & Pain Points")
field(
    "What is the primary risk or problem that brought you here?",
    hint="e.g. no audit trail for agent actions, credential sprawl, compliance gap, customer ask, board mandate",
    lines=3,
)
field_pair(
    "Has an agent ever taken an action it shouldn't have?",
    "Is this a board / C-suite / legal priority, or a team-level initiative?",
    hint_a="Yes / No — brief description if yes",
    hint_b="Who is the internal sponsor for this?",
)
field(
    "What does 'good' look like at the end of a 90-day pilot?",
    hint="Be specific — e.g. 'every agent action is logged and attributable', 'we can produce a SOC 2 evidence pack on demand'",
    lines=3,
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 05 — Tech Stack & Infrastructure
# ══════════════════════════════════════════════════════════════════════════

section_header("05", "Tech Stack & Infrastructure")
field_pair(
    "Which AI / agent frameworks are you using?",
    "Where is your infrastructure hosted?",
    hint_a="e.g. LangChain, CrewAI, custom, OpenAI Assistants",
    hint_b="e.g. AWS, GCP, Azure, hybrid",
)
field_pair(
    "Identity provider (IdP) in use?",
    "SIEM or log-aggregation tooling?",
    hint_a="e.g. Okta, Azure AD, Google Workspace",
    hint_b="e.g. Splunk, Datadog, none yet",
)
field(
    "Any hard integration requirements or constraints we should know about?",
    hint="e.g. must integrate with existing SIEM, no new SaaS vendors without security review",
    lines=2,
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 06 — Stakeholders & Timeline
# ══════════════════════════════════════════════════════════════════════════

section_header("06", "Stakeholders & Timeline")
field(
    "Who else needs to be involved in the decision?",
    hint="e.g. CISO, GC, CTO, procurement — name / role / what they care about",
    lines=2,
)
field_pair(
    "What is your target start date for a pilot?",
    "Is there a hard deadline driving urgency?",
    hint_a="e.g. Q3 2026, before next audit",
    hint_b="e.g. contract renewal, board review, audit date",
)
field_pair(
    "Have you evaluated any alternatives?",
    "Approximate annual budget range for this category",
    hint_a="e.g. build in-house, competitor, doing nothing",
    hint_b="e.g. <$50K, $50–150K, $150K+, unknown",
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 07 — Anything Else
# ══════════════════════════════════════════════════════════════════════════

section_header("07", "Anything Else We Should Know?", "Open — add whatever didn't fit above")
field("Open notes — context, concerns, questions you want to make sure we cover", lines=4)

# ══════════════════════════════════════════════════════════════════════════
# JEFF'S NOTES — internal section
# ══════════════════════════════════════════════════════════════════════════

spacer(6)
# Divider
dp = doc.add_paragraph()
set_para_spacing(dp, before=0, after=0)
r = dp.add_run("─" * 80)
r.font.color.rgb = LINE
r.font.size = Pt(6)

spacer(4)
notes_block("JEFF'S NOTES  —  Internal use only", lines=8)


# ── Save ───────────────────────────────────────────────────────────────────
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Saved: {OUT}")
